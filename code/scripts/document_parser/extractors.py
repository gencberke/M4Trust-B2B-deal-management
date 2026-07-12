"""Concrete TextExtractor implementations, one per source format.

DigitalPdfExtractor and OcrPdfExtractor are single-purpose. HybridPdfExtractor
composes them (digital text first, OCR fallback per page) so callers never
have to know or check whether a given PDF is a scan -- mixed documents
(some digital pages, some scanned) are handled transparently.
"""

import logging
import shutil
from importlib.metadata import PackageNotFoundError, version
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
import pytesseract
from docx import Document
from PIL import Image
from pytesseract import Output

from .exceptions import ExtractionError
from .interfaces import TextExtractor

logger = logging.getLogger(__name__)

# Below this many characters of digital text, a PDF page is treated as a
# scan (image-only) rather than as a page with a real text layer.
MIN_DIGITAL_CHARS_PER_PAGE = 20

# The Windows installer sometimes skips "add to PATH" even when the binary
# itself is installed. Fall back to the well-known default install location
# so OCR still works without requiring a manual PATH edit on every machine.
_WINDOWS_DEFAULT_TESSERACT = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
if not shutil.which("tesseract") and Path(_WINDOWS_DEFAULT_TESSERACT).exists():
    pytesseract.pytesseract.tesseract_cmd = _WINDOWS_DEFAULT_TESSERACT
    logger.info("tesseract not on PATH, using default install path %s", _WINDOWS_DEFAULT_TESSERACT)
OCR_RENDER_DPI = 300
OCR_LANGUAGE = "tur"


def _package_version(name: str) -> str | None:
    try:
        return version(name)
    except PackageNotFoundError:
        return None


def _tesseract_version() -> str | None:
    try:
        return str(pytesseract.get_tesseract_version()).splitlines()[0]
    except Exception:
        return None


class DigitalPdfExtractor(TextExtractor):
    """Extracts text from PDF pages that already carry a text layer."""

    def extract_page(self, page: "fitz.Page") -> str:
        return page.get_text().strip()

    def extract(self, file_path: Path) -> str:
        try:
            with fitz.open(file_path) as doc:
                return "\n\n".join(self.extract_page(page) for page in doc)
        except Exception as exc:
            raise ExtractionError("Digital PDF extraction failed.") from exc


class OcrPdfExtractor(TextExtractor):
    """Extracts text from PDF pages via Tesseract, rendering each page to an image first.

    Requires the Tesseract binary to be installed on the host machine --
    pytesseract is only a wrapper around that executable, not an OCR engine
    itself.
    """

    def __init__(self, dpi: int = OCR_RENDER_DPI, lang: str = OCR_LANGUAGE):
        self.dpi = dpi
        self.lang = lang
        self.last_confidence: float | None = None

    def extract_page(self, page: "fitz.Page") -> str:
        pixmap = page.get_pixmap(dpi=self.dpi)
        image = Image.frombytes("RGB", (pixmap.width, pixmap.height), pixmap.samples)
        data = pytesseract.image_to_data(
            image, lang=self.lang, output_type=Output.DICT
        )
        lines: dict[tuple[int, int, int], list[str]] = {}
        confidences: list[float] = []
        for index, raw_text in enumerate(data.get("text", [])):
            text = str(raw_text).strip()
            if not text:
                continue
            key = (
                int(data["block_num"][index]),
                int(data["par_num"][index]),
                int(data["line_num"][index]),
            )
            lines.setdefault(key, []).append(text)
            try:
                confidence = float(data["conf"][index])
            except (TypeError, ValueError):
                continue
            if confidence >= 0:
                confidences.append(confidence / 100.0)
        self.last_confidence = (
            sum(confidences) / len(confidences) if confidences else None
        )
        return "\n".join(" ".join(words) for words in lines.values()).strip()

    def extract(self, file_path: Path) -> str:
        try:
            with fitz.open(file_path) as doc:
                return "\n\n".join(self.extract_page(page) for page in doc)
        except pytesseract.TesseractNotFoundError as exc:
            raise ExtractionError(
                "Tesseract OCR engine not found on this machine. It is a "
                "system binary, not a pip package -- install it separately "
                "and make sure it is on PATH."
            ) from exc
        except Exception as exc:
            raise ExtractionError("OCR extraction failed.") from exc


class HybridPdfExtractor(TextExtractor):
    """The PDF extractor registered with the factory.

    Tries digital text extraction per page; any page that comes back
    (near-)empty is assumed to be a scanned image and is re-extracted via
    OCR, so mixed PDFs are handled without the caller having to know which
    pages are which.
    """

    def __init__(
        self,
        digital_extractor: Optional[DigitalPdfExtractor] = None,
        ocr_extractor: Optional[OcrPdfExtractor] = None,
        min_digital_chars: int = MIN_DIGITAL_CHARS_PER_PAGE,
    ):
        self.digital_extractor = digital_extractor or DigitalPdfExtractor()
        self.ocr_extractor = ocr_extractor or OcrPdfExtractor()
        self.min_digital_chars = min_digital_chars
        self.last_provenance: dict = {}

    def extract(self, file_path: Path) -> str:
        try:
            with fitz.open(file_path) as doc:
                pages = []
                ocr_pages: list[int] = []
                ocr_confidences: list[float] = []
                for i, page in enumerate(doc):
                    text = self.digital_extractor.extract_page(page)
                    if len(text) < self.min_digital_chars:
                        logger.info(
                            "Page %d has no digital text layer; OCR fallback",
                            i,
                        )
                        text = self.ocr_extractor.extract_page(page)
                        ocr_pages.append(i + 1)
                        confidence = getattr(self.ocr_extractor, "last_confidence", None)
                        if isinstance(confidence, (int, float)):
                            ocr_confidences.append(float(confidence))
                    pages.append(text)
                self.last_provenance = {
                    "document_engine": "pymupdf",
                    "document_engine_version": getattr(fitz, "VersionBind", None),
                    "ocr_engine": "tesseract" if ocr_pages else None,
                    "ocr_version": _tesseract_version() if ocr_pages else None,
                    "ocr_confidence": (
                        sum(ocr_confidences) / len(ocr_confidences)
                        if ocr_confidences
                        else None
                    ),
                    "page_count": len(doc),
                    "ocr_pages": ocr_pages,
                }
                return "\n\n".join(pages)
        except ExtractionError:
            raise
        except Exception as exc:
            raise ExtractionError("PDF extraction failed.") from exc


class DocxExtractor(TextExtractor):
    """Extracts text from DOCX files, including paragraph and table content."""

    def extract(self, file_path: Path) -> str:
        try:
            document = Document(file_path)
        except Exception as exc:
            raise ExtractionError("DOCX extraction failed.") from exc

        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        self.last_provenance = {
            "document_engine": "python-docx",
            "document_engine_version": _package_version("python-docx"),
            "ocr_engine": None,
            "ocr_version": None,
            "ocr_confidence": None,
            "page_count": None,
            "ocr_pages": [],
        }
        return "\n\n".join(parts)
